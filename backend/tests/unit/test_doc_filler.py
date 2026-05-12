"""Unit tests for app/services/doc_filler.py and the new forms endpoints.

All tests build in-memory .docx documents using python-docx and redirect
FORM_SOURCES_DIR via monkeypatch — no real form source files are required.
_convert_docx_to_pdf is mocked in all fill_doc tests so LibreOffice is not
required in the test environment.

Tests:
  test_fill_doc_replaces_dot_sequence
  test_fill_doc_skips_signing_section
  test_fill_doc_fills_cccd_grid
  test_fill_doc_prefills_family_table_row1
  test_fill_endpoint_invalid_form_file_returns_422
  test_configs_endpoint_tthc001_returns_empty_forms
  test_configs_endpoint_returns_fields_for_valid_procedure
  test_fill_endpoint_returns_pdf_bytes
  test_fill_doc_missing_field_value_leaves_dots
  test_fill_doc_signing_section_in_table_not_modified
"""

from __future__ import annotations

import io
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient

from app.main import app

# Fake PDF bytes returned by the _convert_docx_to_pdf mock in all tests.
_FAKE_PDF = b"%PDF-1.4 fake"


# ---------------------------------------------------------------------------
# Fixtures — build temporary .docx files
# ---------------------------------------------------------------------------

def _save_doc(doc: Document, path: Path) -> None:
    """Save a python-docx Document to *path*."""
    doc.save(str(path))


@pytest.fixture
def form_dir(tmp_path, monkeypatch):
    """Redirect FORM_SOURCES_DIR to *tmp_path* for isolation."""
    import app.services.doc_filler as df

    monkeypatch.setattr(df, "FORM_SOURCES_DIR", tmp_path)
    return tmp_path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_simple_doc(tmp_path: Path, filename: str, para_text: str) -> Path:
    """Create a .docx at tmp_path/filename with one paragraph."""
    doc = Document()
    doc.add_paragraph(para_text)
    path = tmp_path / filename
    _save_doc(doc, path)
    return path


def _make_table_doc(
    tmp_path: Path,
    filename: str,
    rows: int,
    cols: int,
    cell_texts: list[list[str]] | None = None,
    bold_header: bool = False,
) -> Path:
    """Create a .docx at tmp_path/filename with one table."""
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    if cell_texts:
        for r, row_texts in enumerate(cell_texts):
            for c, text in enumerate(row_texts):
                cell = table.rows[r].cells[c]
                para = cell.paragraphs[0]
                run = para.add_run(text)
                if bold_header and r == 0:
                    run.bold = True
    path = tmp_path / filename
    _save_doc(doc, path)
    return path


# ---------------------------------------------------------------------------
# Test 1 — dot sequence replacement
# ---------------------------------------------------------------------------

async def test_fill_doc_replaces_dot_sequence(form_dir, monkeypatch):
    """Dot sequences in a paragraph are replaced with the matching field value."""
    from app.services.doc_filler import fill_doc

    # Create a form doc with a labelled dot sequence
    doc = Document()
    doc.add_paragraph("Họ, chữ đệm và tên: ...................")
    path = form_dir / "1.MuCT01banhnhkmtheoThngts53.doc"
    _save_doc(doc, path)

    # _convert_docx_to_pdf receives the filled .docx bytes; capture them for
    # inspection by having the mock return them as-is (still a valid .docx).
    captured: list[bytes] = []

    async def _capture(docx_bytes: bytes) -> bytes:
        captured.append(docx_bytes)
        return _FAKE_PDF

    with patch(
        "app.services.doc_filler._convert_docx_to_pdf",
        side_effect=_capture,
    ):
        result_bytes = await fill_doc(
            "1.MuCT01banhnhkmtheoThngts53.doc",
            {"ho_chu_dem_va_ten": "Nguyễn Văn A"},
        )

    assert result_bytes[:4] == b"%PDF", "fill_doc must return PDF bytes"
    # Verify the substitution happened in the intermediate .docx
    result_doc = Document(io.BytesIO(captured[0]))
    full_text = " ".join(p.text for p in result_doc.paragraphs)
    assert "Nguyễn Văn A" in full_text, (
        f"Expected 'Nguyễn Văn A' in filled document, got: {full_text!r}"
    )


# ---------------------------------------------------------------------------
# Test 2 — signing sections are not modified
# ---------------------------------------------------------------------------

async def test_fill_doc_skips_signing_section(form_dir, monkeypatch):
    """Paragraphs containing signing phrases must not be modified."""
    from app.services.doc_filler import fill_doc

    original_sign_text = "Ký, ghi rõ họ tên ..........."
    doc = Document()
    doc.add_paragraph("Họ, chữ đệm và tên: ...................")
    doc.add_paragraph(original_sign_text)
    path = form_dir / "1.MuCT01banhnhkmtheoThngts53.doc"
    _save_doc(doc, path)

    captured: list[bytes] = []

    async def _capture(docx_bytes: bytes) -> bytes:
        captured.append(docx_bytes)
        return _FAKE_PDF

    with patch(
        "app.services.doc_filler._convert_docx_to_pdf",
        side_effect=_capture,
    ):
        result_bytes = await fill_doc(
            "1.MuCT01banhnhkmtheoThngts53.doc",
            {"ho_chu_dem_va_ten": "Nguyễn Văn A"},
        )

    assert result_bytes[:4] == b"%PDF", "fill_doc must return PDF bytes"
    result_doc = Document(io.BytesIO(captured[0]))
    paragraphs = [p.text for p in result_doc.paragraphs]

    # The signing paragraph must be present and unchanged (dots not replaced)
    sign_paragraphs = [p for p in paragraphs if "Ký, ghi rõ họ tên" in p]
    assert sign_paragraphs, "Signing paragraph disappeared from output"
    assert "Nguyễn Văn A" not in sign_paragraphs[0], (
        f"Signing section was incorrectly modified: {sign_paragraphs[0]!r}"
    )


# ---------------------------------------------------------------------------
# Test 3 — CCCD character grid fill
# ---------------------------------------------------------------------------

async def test_fill_doc_fills_cccd_grid(form_dir, monkeypatch):
    """A 1-row, 12-cell grid table is filled character by character from id_number."""
    from app.services.doc_filler import fill_doc

    id_number = "012345678901"

    # Build a doc with a 1-row, 12-cell table (CCCD grid)
    doc = Document()
    table = doc.add_table(rows=1, cols=12)
    for cell in table.rows[0].cells:
        cell.paragraphs[0].add_run(".")
    path = form_dir / "1.MuCT01banhnhkmtheoThngts53.doc"
    _save_doc(doc, path)

    captured: list[bytes] = []

    async def _capture(docx_bytes: bytes) -> bytes:
        captured.append(docx_bytes)
        return _FAKE_PDF

    with patch(
        "app.services.doc_filler._convert_docx_to_pdf",
        side_effect=_capture,
    ):
        result_bytes = await fill_doc(
            "1.MuCT01banhnhkmtheoThngts53.doc",
            {"so_dinh_danh_ca_nhan": id_number},
        )

    assert result_bytes[:4] == b"%PDF", "fill_doc must return PDF bytes"
    result_doc = Document(io.BytesIO(captured[0]))
    cells = result_doc.tables[0].rows[0].cells
    grid_text = "".join(c.text for c in cells)
    assert grid_text == id_number, (
        f"CCCD grid content mismatch. Expected {id_number!r}, got {grid_text!r}"
    )


# ---------------------------------------------------------------------------
# Test 4 — family member table row 1 pre-fill
# ---------------------------------------------------------------------------

async def test_fill_doc_prefills_family_table_row1(form_dir, monkeypatch):
    """Row 1 of a family member table is pre-filled; rows 2+ are untouched."""
    from app.services.doc_filler import fill_doc

    # Build a doc with a family table:
    #   Row 0 (header): bold column names
    #   Row 1 (data):   empty cells to be filled
    #   Row 2 (data):   empty cells — must remain untouched
    doc = Document()
    table = doc.add_table(rows=3, cols=3)

    # Header row — bold labels (specific enough to uniquely match each field
    # via word-subset: "cha nuôi" / "con nuôi" disambiguate from other name
    # and date fields in the same config)
    headers = ["Họ tên cha nuôi", "Ngày sinh cha nuôi", "Giới tính con nuôi"]
    for c, label in enumerate(headers):
        para = table.rows[0].cells[c].paragraphs[0]
        run = para.add_run(label)
        run.bold = True

    path = form_dir / "7.Tkhaingklivicnuiconnui.doc"
    _save_doc(doc, path)

    captured: list[bytes] = []

    async def _capture(docx_bytes: bytes) -> bytes:
        captured.append(docx_bytes)
        return _FAKE_PDF

    with patch(
        "app.services.doc_filler._convert_docx_to_pdf",
        side_effect=_capture,
    ):
        result_bytes = await fill_doc(
            "7.Tkhaingklivicnuiconnui.doc",
            {
                "ho_ten_cha_nuoi": "Nguyễn Văn A",
                "ngay_sinh_cha_nuoi": "01/01/1980",
                "gioi_tinh_con_nuoi": "Nam",
            },
        )

    assert result_bytes[:4] == b"%PDF", "fill_doc must return PDF bytes"
    result_doc = Document(io.BytesIO(captured[0]))
    tbl = result_doc.tables[0]

    # Row 1 should have been filled for "Họ tên cha nuôi" column (index 0)
    row1_texts = [tbl.rows[1].cells[c].text for c in range(3)]
    assert any("Nguyễn Văn A" in t for t in row1_texts), (
        f"Row 1 not pre-filled with name. Row 1 texts: {row1_texts}"
    )

    # Row 2 must remain empty (untouched)
    row2_texts = [tbl.rows[2].cells[c].text for c in range(3)]
    assert all(t.strip() == "" for t in row2_texts), (
        f"Row 2 was unexpectedly modified: {row2_texts}"
    )


# ---------------------------------------------------------------------------
# Test 5 — fill endpoint returns 422 when form_file not valid for procedure_id
# ---------------------------------------------------------------------------

async def test_fill_endpoint_invalid_form_file_returns_422():
    """POST /api/v1/forms/fill returns 422 when form_file not in procedure's list."""
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        resp = await client.post(
            "/api/v1/forms/fill",
            json={
                "procedure_id": "TTHC-CR-001",
                "form_file": "7.Tkhaingklivicnuiconnui.doc",  # belongs to TTHC-AD-002
                "field_values": {},
            },
        )

    assert resp.status_code == 422, (
        f"Expected 422 for invalid form_file, got {resp.status_code}: {resp.text}"
    )
    assert "form_file not valid for procedure_id" in resp.json()["detail"]


# ---------------------------------------------------------------------------
# Test 6 — configs endpoint returns empty forms for TTHC-001
# ---------------------------------------------------------------------------

async def test_configs_endpoint_tthc001_returns_empty_forms():
    """GET /api/v1/forms/configs/TTHC-001 returns an empty forms array."""
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        resp = await client.get("/api/v1/forms/configs/TTHC-001")

    assert resp.status_code == 200, (
        f"Expected 200 for TTHC-001, got {resp.status_code}: {resp.text}"
    )
    data = resp.json()
    assert data["procedure_id"] == "TTHC-001"
    assert data["forms"] == [], (
        f"Expected empty forms for TTHC-001, got: {data['forms']}"
    )


# ---------------------------------------------------------------------------
# Test 7 — configs endpoint returns fields for a valid procedure
# ---------------------------------------------------------------------------

async def test_configs_endpoint_returns_fields_for_valid_procedure():
    """GET /api/v1/forms/configs/TTHC-CR-001 returns two form tabs with fields."""
    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        resp = await client.get("/api/v1/forms/configs/TTHC-CR-001")

    assert resp.status_code == 200, (
        f"Expected 200, got {resp.status_code}: {resp.text}"
    )
    data = resp.json()
    assert data["procedure_id"] == "TTHC-CR-001"
    assert len(data["forms"]) == 2, (
        f"Expected 2 form tabs for TTHC-CR-001, got {len(data['forms'])}"
    )
    # Each form tab must have required keys
    for form in data["forms"]:
        assert "form_file" in form
        assert "tab_label" in form
        assert "fields" in form
        assert isinstance(form["fields"], list)
        assert len(form["fields"]) > 0


# ---------------------------------------------------------------------------
# Test 8 — fill endpoint returns docx bytes on success
# ---------------------------------------------------------------------------

async def test_fill_endpoint_returns_pdf_bytes(form_dir, monkeypatch):
    """POST /api/v1/forms/fill returns a PDF file with correct Content-Type."""
    # Create a real .docx at the expected path (PROCEDURE_FORM_FILES uses .docx extension)
    doc = Document()
    doc.add_paragraph("Họ, chữ đệm và tên: ...................")
    path = form_dir / "1.MuCT01banhnhkmtheoThngts53.docx"
    _save_doc(doc, path)

    # Monkeypatch FORM_SOURCES_DIR inside the running app module
    import app.services.doc_filler as df
    monkeypatch.setattr(df, "FORM_SOURCES_DIR", form_dir)

    with patch(
        "app.services.doc_filler._convert_docx_to_pdf",
        new_callable=AsyncMock,
        return_value=_FAKE_PDF,
    ):
        async with AsyncClient(
            transport=ASGITransport(app=app), base_url="http://test"
        ) as client:
            resp = await client.post(
                "/api/v1/forms/fill",
                json={
                    "procedure_id": "TTHC-002",
                    "form_file": "1.MuCT01banhnhkmtheoThngts53.docx",
                    "field_values": {"ho_chu_dem_va_ten": "Nguyễn Văn A"},
                },
            )

    assert resp.status_code == 200, (
        f"Expected 200, got {resp.status_code}: {resp.text}"
    )
    assert resp.headers.get("content-type") == "application/pdf", (
        f"Unexpected Content-Type: {resp.headers.get('content-type')}"
    )
    assert "attachment" in resp.headers.get("content-disposition", ""), (
        f"Missing Content-Disposition: {resp.headers.get('content-disposition')}"
    )
    assert ".pdf" in resp.headers.get("content-disposition", ""), (
        f"Expected .pdf in Content-Disposition: {resp.headers.get('content-disposition')}"
    )
    assert resp.content[:4] == b"%PDF", "Response is not a valid PDF (missing %PDF magic bytes)"


# ---------------------------------------------------------------------------
# Test 9 — missing field value leaves dots unchanged
# ---------------------------------------------------------------------------

async def test_fill_doc_missing_field_value_leaves_dots(form_dir, monkeypatch):
    """When a field has no value in field_values, the dots remain unchanged."""
    from app.services.doc_filler import fill_doc

    original_dots = "Họ, chữ đệm và tên: ..................."
    doc = Document()
    doc.add_paragraph(original_dots)
    path = form_dir / "1.MuCT01banhnhkmtheoThngts53.doc"
    _save_doc(doc, path)

    captured: list[bytes] = []

    async def _capture(docx_bytes: bytes) -> bytes:
        captured.append(docx_bytes)
        return _FAKE_PDF

    with patch(
        "app.services.doc_filler._convert_docx_to_pdf",
        side_effect=_capture,
    ):
        # Pass field_values without the matching key
        result_bytes = await fill_doc(
            "1.MuCT01banhnhkmtheoThngts53.doc",
            {"ngay_thang_nam_sinh": "01/01/1990"},  # unrelated field
        )

    assert result_bytes[:4] == b"%PDF", "fill_doc must return PDF bytes"
    result_doc = Document(io.BytesIO(captured[0]))
    full_text = " ".join(p.text for p in result_doc.paragraphs)
    # Dots must still be present
    import re
    assert re.search(r"\.{3,}", full_text), (
        f"Expected dots to remain when field value missing, but got: {full_text!r}"
    )


# ---------------------------------------------------------------------------
# Test 10 — signing section inside a table cell is not modified
# ---------------------------------------------------------------------------

async def test_fill_doc_signing_section_in_table_not_modified(form_dir, monkeypatch):
    """A table cell containing a signing phrase must not have its dots replaced."""
    from app.services.doc_filler import fill_doc

    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    # Row 0: normal data cell
    table.rows[0].cells[0].paragraphs[0].add_run("Họ, chữ đệm và tên: ...........")
    table.rows[0].cells[1].paragraphs[0].add_run("")

    # Row 1: signing section cell — must not be modified
    signing_text = "Ký, ghi rõ họ tên ..........."
    table.rows[1].cells[0].paragraphs[0].add_run(signing_text)
    table.rows[1].cells[1].paragraphs[0].add_run("")

    path = form_dir / "1.MuCT01banhnhkmtheoThngts53.doc"
    _save_doc(doc, path)

    captured: list[bytes] = []

    async def _capture(docx_bytes: bytes) -> bytes:
        captured.append(docx_bytes)
        return _FAKE_PDF

    with patch(
        "app.services.doc_filler._convert_docx_to_pdf",
        side_effect=_capture,
    ):
        result_bytes = await fill_doc(
            "1.MuCT01banhnhkmtheoThngts53.doc",
            {"ho_chu_dem_va_ten": "Nguyễn Văn A"},
        )

    assert result_bytes[:4] == b"%PDF", "fill_doc must return PDF bytes"
    result_doc = Document(io.BytesIO(captured[0]))
    tbl = result_doc.tables[0]
    sign_cell_text = tbl.rows[1].cells[0].text

    assert "Ký, ghi rõ họ tên" in sign_cell_text, "Signing phrase was removed"
    assert "Nguyễn Văn A" not in sign_cell_text, (
        f"Signing section was incorrectly filled: {sign_cell_text!r}"
    )
