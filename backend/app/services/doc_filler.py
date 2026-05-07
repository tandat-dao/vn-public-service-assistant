"""DOC form filling service.

Opens a .doc (OOXML-based) file from FORM_SOURCES_DIR, applies field value
substitutions using four rules, and returns the filled document as bytes.

Fill rules applied in order:
  Rule 1 — Dot sequence replacement in paragraphs and table cells.
             Signing sections are skipped entirely.
  Rule 2 — CCCD character grid fill (1-row, 9+ cell tables).
  Rule 3 — Family member table row 1 pre-fill.
  Rule 4 — Signing sections are never modified under any circumstance.

All operations are best-effort: missing field values leave the document
unchanged for that position. Never raises on missing field values.
"""

from __future__ import annotations

import asyncio
import io
import os
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.core.form_field_configs import FORM_FILE_CONFIGS, FormField

FORM_SOURCES_DIR = Path(__file__).parent.parent.parent / "data" / "form_sources"

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Matches 3+ consecutive dots, middle dots (·), horizontal ellipsis (…),
# or mathematical ellipsis (⋯).
_DOT_RE = re.compile(r"[.·…⋯]{3,}")

# Phrases that mark a signing section — never modify paragraphs/cells
# containing any of these.
_SIGNING_PHRASES: tuple[str, ...] = (
    "Ký, ghi rõ họ tên",
    "Xác nhận của",
    "ngày....tháng",
    "CHỮ KÝ",
    "NGƯỜI KÝ",
)

# Column header keywords that indicate a family-member table row.
_FAMILY_COL_KEYWORDS: tuple[str, ...] = (
    "họ",
    "tên",
    "sinh",
    "giới",
    "tính",
    "định danh",
    "quốc tịch",
    "dân tộc",
    "quan hệ",
)


# ---------------------------------------------------------------------------
# Helpers — signing-section detection
# ---------------------------------------------------------------------------

def _is_signing_section(text: str) -> bool:
    """Return True when *text* contains any signing-section phrase."""
    return any(phrase in text for phrase in _SIGNING_PHRASES)


# ---------------------------------------------------------------------------
# Helpers — label-to-field matching
# ---------------------------------------------------------------------------

def _normalize(text: str) -> str:
    """Lowercase, strip dots/colons/parens/footnote markers for label comparison."""
    text = re.sub(r"\(\d+\)", "", text)   # strip footnote refs like (1), (2), (5)
    text = _DOT_RE.sub("", text)
    text = re.sub(r"[,:()\[\]/\\]+", " ", text)
    return text.strip().lower()


def _find_field_id(
    label: str,
    fields: list[FormField],
    consumed: set[str] | None = None,
) -> str | None:
    """Return the field ID whose label best matches *label*, skipping consumed IDs.

    Matching priority:
      1. Substring: either label is a substring of the other (exact).
      2. Word-subset: all words in the shorter token set appear in the longer.

    *consumed* tracks field IDs already filled so repeated labels (e.g. "Năm sinh"
    for mother then father) map to distinct fields in document order.
    """
    norm = _normalize(label)
    if not norm:
        return None

    # Pass 1 — substring match
    for field in fields:
        if consumed and field["id"] in consumed:
            continue
        norm_field = _normalize(field["label"])
        if norm_field and (norm_field in norm or norm in norm_field):
            return field["id"]

    # Pass 2 — word-subset match
    query_words = set(norm.split())
    if not query_words:
        return None
    for field in fields:
        if consumed and field["id"] in consumed:
            continue
        norm_field = _normalize(field["label"])
        if not norm_field:
            continue
        field_words = set(norm_field.split())
        shorter, longer = (
            (query_words, field_words)
            if len(query_words) <= len(field_words)
            else (field_words, query_words)
        )
        if shorter and shorter.issubset(longer):
            return field["id"]

    return None


# ---------------------------------------------------------------------------
# Helpers — cell text writing
# ---------------------------------------------------------------------------

def _set_cell_char(cell, char: str) -> None:
    """Write a single character into the first run of the cell's first paragraph."""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = char
    else:
        para.add_run(char)


def _set_cell_value(cell, value: str) -> None:
    """Replace the text content of the first paragraph in *cell* with *value*."""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    # Clear all existing runs
    for run in para.runs:
        run.text = ""
    # Write value into the first run (or add one if none exist)
    if para.runs:
        para.runs[0].text = value
    else:
        para.add_run(value)


# ---------------------------------------------------------------------------
# Rule 1 — Paragraph field filling (tab-based and dot-based)
# ---------------------------------------------------------------------------

def _merge_dot_runs(para) -> None:
    """Merge consecutive runs that are entirely dots into the first one.

    Word sometimes splits a single dot sequence across multiple runs. Merging
    them lets _apply_paragraph_fields see the full sequence in one run and
    fill it with a single value rather than leaving subsequent dot runs unfilled.
    """
    runs = para.runs
    i = 0
    while i < len(runs) - 1:
        curr_text = runs[i].text or ""
        next_text = runs[i + 1].text or ""
        if curr_text and next_text and _DOT_RE.fullmatch(curr_text) and _DOT_RE.fullmatch(next_text):
            runs[i].text = curr_text + next_text
            runs[i + 1].text = ""
        else:
            i += 1


def _apply_paragraph_fields(
    para,
    label_hint: str,
    fields: list[FormField],
    field_values: dict[str, str],
    consumed: set[str],
) -> str:
    """Fill one paragraph and return the label text to carry into the next.

    Handles three patterns found in Vietnamese government forms:

    1. Tab-based  — "Label:\t"  → value inserted after the tab.
    2. Dot-based  — "Label:....." → dots replaced by value.
    3. Multi-field — "F1:..... F2:..... F3:\t" → left-to-right label tracking.

    *consumed* is updated in-place whenever a field is successfully filled so
    that repeated labels (e.g. "Năm sinh" for mother then father) resolve to
    distinct field IDs in document order.
    """
    if _is_signing_section(para.text):
        return para.text

    has_dots = _DOT_RE.search(para.text)
    has_tab = "\t" in para.text
    if not has_dots and not has_tab:
        return para.text  # nothing to fill; pass text as hint to next paragraph

    _merge_dot_runs(para)

    label_acc = label_hint  # accumulates label text as we scan runs left→right

    for run in para.runs:
        text = run.text
        if not text:
            continue
        if _is_signing_section(text):
            break

        if _DOT_RE.search(text):
            parts = _DOT_RE.split(text)       # text segments between dot groups
            dot_groups = _DOT_RE.findall(text) # the dot groups themselves
            new_text = ""

            for i, dot_group in enumerate(dot_groups):
                before = parts[i]
                label = (label_acc + before).strip()
                field_id = _find_field_id(label, fields, consumed) if label else None
                value = field_values.get(field_id) if field_id else None
                if field_id and value:
                    consumed.add(field_id)
                    # Add a space before the value when the accumulated label ends
                    # without one (e.g. "Giới tính:" → "Giới tính: Nam").
                    # When before="" the value is at the start of this run but may
                    # directly follow text from the previous run — always add a space.
                    prefix = " " if not before or not before[-1].isspace() else ""
                    new_text += before + prefix + value
                else:
                    new_text += before + dot_group
                label_acc = ""  # reset after each dot group

            # Text after the last dot group — may contain a tab for another field.
            # Add a space if the filled value runs directly into the next label text.
            last = parts[-1]
            if new_text and not new_text[-1].isspace() and last and not last[0].isspace():
                new_text += " "

            # When this run ends with a filled value and the *next* non-empty run
            # starts with a non-space character, append a trailing space so runs
            # from different XML elements don't collide (e.g. "1992Dân tộc:").
            if new_text and not new_text[-1].isspace() and not last:
                found = False
                for nr in para.runs:
                    if not found:
                        if nr._element is run._element:
                            found = True
                        continue
                    if nr.text:
                        if not nr.text[0].isspace():
                            new_text += " "
                        break
            if "\t" in last:
                tab_i = last.index("\t")
                label = (label_acc + last[:tab_i]).strip()
                field_id = _find_field_id(label, fields, consumed) if label else None
                value = field_values.get(field_id) if field_id else None
                if field_id and value:
                    consumed.add(field_id)
                    # Replace the tab with a space so tab-leader dots don't appear.
                    new_text += last[:tab_i].rstrip() + " " + value
                    label_acc = ""
                else:
                    new_text += last
                    label_acc = last[tab_i:]
            else:
                new_text += last
                label_acc = last

            run.text = new_text

        elif "\t" in text:
            tab_i = text.index("\t")
            label = (label_acc + text[:tab_i]).strip()
            field_id = _find_field_id(label, fields, consumed) if label else None
            value = field_values.get(field_id) if field_id else None
            if field_id and value:
                consumed.add(field_id)
                # Replace the tab with a space so tab-leader dots don't appear.
                run.text = text[:tab_i].rstrip() + " " + value
                label_acc = ""
            else:
                label_acc = text[tab_i:]

        else:
            label_acc += text

    return label_acc


# ---------------------------------------------------------------------------
# Rule 2 — CCCD character grid detection and fill
# ---------------------------------------------------------------------------

def _is_cccd_grid(table) -> bool:
    """True when the table looks like a 12-cell CCCD character grid."""
    rows = table.rows
    if len(rows) != 1:
        return False
    cells = rows[0].cells
    if len(cells) < 9:
        return False
    for cell in cells:
        text = cell.text.strip()
        if len(text) > 1:
            return False
    return True


def _find_id_number(field_values: dict[str, str]) -> str | None:
    """Return the first value in *field_values* that looks like a CCCD number."""
    _ID_RE = re.compile(r"^\d{9,12}$")
    for value in field_values.values():
        if value and _ID_RE.match(value.strip()):
            return value.strip()
    return None


def _apply_cccd_grid_rule(table, field_values: dict[str, str]) -> None:
    """Fill the CCCD character grid one digit per cell."""
    id_number = _find_id_number(field_values)
    if not id_number:
        return
    cells = table.rows[0].cells
    for i, char in enumerate(id_number):
        if i >= len(cells):
            break
        _set_cell_char(cells[i], char)


# ---------------------------------------------------------------------------
# Rule 3 — Family member table row 1 pre-fill
# ---------------------------------------------------------------------------

def _is_header_row(row) -> bool:
    """True when the row looks like a table header (bold text or all-caps labels)."""
    row_text = " ".join(c.text.strip() for c in row.cells)
    if not row_text.strip():
        return False

    # All-caps check with family-relevant keywords
    lower = row_text.lower()
    keyword_hits = sum(1 for kw in _FAMILY_COL_KEYWORDS if kw in lower)
    if keyword_hits >= 2:
        return True

    # Bold text in any cell
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if run.bold and run.text.strip():
                    return True

    return False


def _is_family_table(table) -> bool:
    """True when the table has 2+ rows and a recognisable header row."""
    rows = table.rows
    if len(rows) < 2:
        return False
    return _is_header_row(rows[0])


def _apply_family_table_rule(
    table,
    fields: list[FormField],
    field_values: dict[str, str],
) -> None:
    """Pre-fill row 1 of a family-member table using column header matching."""
    rows = table.rows
    if len(rows) < 2:
        return

    header_cells = rows[0].cells
    data_cells = rows[1].cells

    for i, header_cell in enumerate(header_cells):
        if i >= len(data_cells):
            break
        header_text = header_cell.text.strip()
        if not header_text or _is_signing_section(header_text):
            continue
        field_id = _find_field_id(header_text, fields)
        if not field_id:
            continue
        value = field_values.get(field_id)
        if not value:
            continue
        if not _is_signing_section(data_cells[i].text):
            _set_cell_value(data_cells[i], value)


def _apply_table_cells(
    table,
    fields: list[FormField],
    field_values: dict[str, str],
    consumed: set[str],
) -> None:
    """Apply paragraph-level field filling to every non-signing cell in *table*."""
    for row in table.rows:
        for cell in row.cells:
            if _is_signing_section(cell.text):
                continue
            hint = ""
            for para in cell.paragraphs:
                hint = _apply_paragraph_fields(para, hint, fields, field_values, consumed)


# ---------------------------------------------------------------------------
# LibreOffice PDF conversion
# ---------------------------------------------------------------------------

# Matches the path used in ingest_full_documents.py
_LIBREOFFICE_EXE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")


def _run_libreoffice(docx_path: Path, tmp_dir_path: Path) -> subprocess.CompletedProcess:
    """Synchronous LibreOffice subprocess call — run via executor to avoid blocking."""
    return subprocess.run(
        [
            str(_LIBREOFFICE_EXE),
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(tmp_dir_path),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
        timeout=60,
    )


async def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert .docx bytes to .pdf bytes via LibreOffice headless.

    Writes docx_bytes to a temp file, invokes LibreOffice to convert
    in a temp directory, reads the resulting .pdf, then cleans up.

    Raises RuntimeError if LibreOffice is not found or conversion fails.
    """
    if not _LIBREOFFICE_EXE.exists():
        raise RuntimeError(
            f"LibreOffice not found at {_LIBREOFFICE_EXE}. "
            "Install LibreOffice or update _LIBREOFFICE_EXE path."
        )

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_dir_path = Path(tmp_dir)
        docx_path = tmp_dir_path / "input.docx"
        docx_path.write_bytes(docx_bytes)

        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(
            None, _run_libreoffice, docx_path, tmp_dir_path
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed (exit {result.returncode}): "
                f"{result.stderr[:500]}"
            )

        pdf_path = tmp_dir_path / "input.pdf"
        if not pdf_path.exists():
            raise RuntimeError(
                f"LibreOffice ran but produced no PDF. stdout: {result.stdout[:300]}"
            )

        return pdf_path.read_bytes()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def fill_doc(
    form_file: str,
    field_values: dict[str, str],
) -> bytes:
    """Open *form_file* from FORM_SOURCES_DIR, apply *field_values*, return bytes.

    Raises FileNotFoundError when form_file is not present in FORM_SOURCES_DIR.
    Raises ValueError when form_file is not registered in FORM_FILE_CONFIGS.
    Fill operations are best-effort — missing values leave the document unchanged.
    The source file on disk is never modified.
    """
    # Normalise: always prefer the .docx version (python-docx can't open binary .doc).
    # This bridges stale in-memory configs that still reference .doc filenames.
    if form_file.endswith(".doc") and not form_file.endswith(".docx"):
        docx_candidate = form_file + "x"
        if (FORM_SOURCES_DIR / docx_candidate).exists():
            form_file = docx_candidate

    config = FORM_FILE_CONFIGS.get(form_file) or FORM_FILE_CONFIGS.get(
        form_file[:-1] if form_file.endswith(".docx") else form_file + "x"
    )
    if config is None:
        raise ValueError(f"form_file not in FORM_FILE_CONFIGS: {form_file!r}")

    source_path = FORM_SOURCES_DIR / form_file
    if not source_path.exists():
        raise FileNotFoundError(f"Form source file not found: {form_file!r}")

    fields: list[FormField] = config["fields"]

    # Load into memory — Document() never modifies the source file on disk.
    with source_path.open("rb") as fh:
        file_bytes = fh.read()

    doc = Document(io.BytesIO(file_bytes))

    consumed: set[str] = set()

    # ── Paragraphs ──────────────────────────────────────────────────────────
    hint = ""
    for para in doc.paragraphs:
        hint = _apply_paragraph_fields(para, hint, fields, field_values, consumed)

    # ── Tables ────────────────────────────────────────────────────────────
    for table in doc.tables:
        if _is_cccd_grid(table):
            _apply_cccd_grid_rule(table, field_values)
            continue
        if _is_family_table(table):
            _apply_family_table_rule(table, fields, field_values)
        _apply_table_cells(table, fields, field_values, consumed)

    # Save the modified document to bytes, convert to PDF, and return.
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    return await _convert_docx_to_pdf(docx_bytes)
